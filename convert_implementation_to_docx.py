#!/usr/bin/env python3
"""
Convert IMPLEMENTATION_SECTION_3_6.md to formatted DOCX with proper styling.
Tables formatted with light blue headers and alternating row colors.
Font: Times New Roman 11pt
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# Color definitions
LIGHT_BLUE_HEADER = "B8CCE4"  # Light blue for header cells
LIGHT_ROW_COLOR = "E7F0F7"    # Light blue for alternating rows
DARK_BLUE = "003366"          # Dark blue for main headings

def set_cell_background(cell, fill_color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def read_markdown_file(filepath):
    """Read the markdown file"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def parse_markdown_content(content):
    """Parse markdown content into structured format"""
    sections = []
    lines = content.split('\n')
    current_section = None
    current_table = None
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Detect headings
        if line.startswith('# '):
            current_section = {'type': 'h1', 'content': line[2:].strip()}
            sections.append(current_section)
        elif line.startswith('## '):
            current_section = {'type': 'h2', 'content': line[3:].strip()}
            sections.append(current_section)
        elif line.startswith('### '):
            current_section = {'type': 'h3', 'content': line[4:].strip()}
            sections.append(current_section)
        elif line.startswith('#### '):
            current_section = {'type': 'h4', 'content': line[5:].strip()}
            sections.append(current_section)
        
        # Detect tables (markdown table starts with |)
        elif line.strip().startswith('|') and not line.strip().startswith('|---'):
            table_rows = []
            # Get header row
            header = [cell.strip() for cell in line.split('|')[1:-1]]
            table_rows.append(header)
            i += 1
            
            # Skip separator row
            if i < len(lines) and lines[i].strip().startswith('|---'):
                i += 1
            
            # Get data rows
            while i < len(lines) and lines[i].strip().startswith('|') and not lines[i].strip().startswith('|---'):
                row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                table_rows.append(row)
                i += 1
            
            sections.append({'type': 'table', 'rows': table_rows})
            i -= 1  # Adjust for the outer loop increment
        
        # Detect horizontal rules
        elif line.strip().startswith('---'):
            sections.append({'type': 'hr'})
        
        # Detect bullet points
        elif line.strip().startswith('- '):
            sections.append({'type': 'bullet', 'content': line.strip()[2:]})
        
        # Regular paragraphs
        elif line.strip() and not line.startswith('|'):
            if not line.startswith('_[') or not line.endswith(']_'):  # Skip placeholder lines as is
                sections.append({'type': 'paragraph', 'content': line.strip()})
            else:
                sections.append({'type': 'placeholder', 'content': line.strip()})
        
        i += 1
    
    return sections

def add_section_to_docx(doc, section):
    """Add a parsed section to the document"""
    if section['type'] == 'h1':
        p = doc.add_heading(section['content'], level=1)
        p_format = p.paragraph_format
        p_format.space_before = Pt(12)
        p_format.space_after = Pt(12)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    elif section['type'] == 'h2':
        p = doc.add_heading(section['content'], level=2)
        p_format = p.paragraph_format
        p_format.space_before = Pt(12)
        p_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    elif section['type'] == 'h3':
        p = doc.add_heading(section['content'], level=3)
        p_format = p.paragraph_format
        p_format.space_before = Pt(10)
        p_format.space_after = Pt(4)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    elif section['type'] == 'h4':
        p = doc.add_heading(section['content'], level=4)
        p_format = p.paragraph_format
        p_format.space_before = Pt(8)
        p_format.space_after = Pt(4)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
    
    elif section['type'] == 'table':
        table = doc.add_table(rows=len(section['rows']), cols=len(section['rows'][0]))
        table.style = 'Light Grid Accent 1'
        
        # Format header row
        for col_idx, cell in enumerate(table.rows[0].cells):
            set_cell_background(cell, LIGHT_BLUE_HEADER)
            cell.text = section['rows'][0][col_idx]
            
            # Format header text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Format data rows with alternating colors
        for row_idx in range(1, len(section['rows'])):
            for col_idx, cell in enumerate(table.rows[row_idx].cells):
                # Alternate row colors
                if row_idx % 2 == 0:
                    set_cell_background(cell, LIGHT_ROW_COLOR)
                
                cell.text = section['rows'][row_idx][col_idx]
                
                # Format cell text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
    
    elif section['type'] == 'hr':
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.space_before = Pt(6)
        p_format.space_after = Pt(6)
    
    elif section['type'] == 'bullet':
        p = doc.add_paragraph(section['content'], style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        p_format = p.paragraph_format
        p_format.space_after = Pt(2)
    
    elif section['type'] == 'placeholder':
        p = doc.add_paragraph(section['content'])
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.italic = True
    
    elif section['type'] == 'paragraph':
        p = doc.add_paragraph(section['content'])
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        p_format = p.paragraph_format
        p_format.space_after = Pt(6)

def create_docx_from_markdown(markdown_path, docx_path):
    """Main function to convert markdown to DOCX"""
    # Read markdown content
    content = read_markdown_file(markdown_path)
    
    # Parse markdown
    sections = parse_markdown_content(content)
    
    # Create document
    doc = Document()
    
    # Set default font for document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Add sections
    for section in sections:
        add_section_to_docx(doc, section)
    
    # Save document
    doc.save(docx_path)
    print(f"✓ DOCX file created: {docx_path}")

if __name__ == "__main__":
    markdown_file = r"c:\Users\awast\OneDrive\Desktop\Sem-4\Student_Placement_Cell_Database_Management_System\IMPLEMENTATION_SECTION_3_6.md"
    docx_file = r"c:\Users\awast\OneDrive\Desktop\Sem-4\Student_Placement_Cell_Database_Management_System\IMPLEMENTATION_SECTION_3_6.docx"
    
    create_docx_from_markdown(markdown_file, docx_file)
