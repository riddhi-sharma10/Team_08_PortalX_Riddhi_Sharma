from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_er_explanation():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('ER DIAGRAM EXPLANATION & RELATIONSHIPS')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Student Placement Cell Database Management System')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Introduction
    intro = doc.add_heading('1. INTRODUCTION', level=1)
    intro.style.font.name = 'Times New Roman'
    
    intro_para = doc.add_paragraph(
        'The Student Placement Cell Database comprises 22 interconnected entities implementing a hierarchical structure that manages the complete placement lifecycle. This document provides a detailed explanation of all entity relationships, cardinalities, and dependencies that form the backbone of the system.'
    )
    intro_para.style.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Hierarchical Overview
    hier = doc.add_heading('2. HIERARCHICAL OVERVIEW', level=1)
    hier.style.font.name = 'Times New Roman'
    
    hier_para = doc.add_paragraph(
        'The database implements a three-tier administrative hierarchy:\n\n'
        'LEVEL 1: CGDC_ADMIN (Career Guidance and Development Cell Administrator)\n'
        'LEVEL 2: PLACEMENT_COORDINATOR (Coordinators managing specific departments)\n'
        'LEVEL 3: STUDENT (End-users seeking placement)\n\n'
        'This hierarchical structure ensures proper access control and management oversight throughout the placement process.'
    )
    hier_para.style.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Key Terminology
    term = doc.add_heading('3. KEY TERMINOLOGY', level=1)
    term.style.font.name = 'Times New Roman'
    
    card_heading = doc.add_heading('Cardinality Notation:', level=2)
    card_heading.style.font.name = 'Times New Roman'
    
    cardinality = [
        '1:1 (One-to-One): One entity instance relates to exactly one instance of another entity',
        '1:N (One-to-Many): One entity instance relates to multiple instances of another entity',
        'M:N (Many-to-Many): Multiple instances relate to multiple instances of another',
        'N:1 (Many-to-One): Reverse of 1:N, indicating the dependent side'
    ]
    
    for item in cardinality:
        p = doc.add_paragraph(item, style='List Bullet')
        p.style.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    dep_heading = doc.add_heading('Participation/Dependency Types:', level=2)
    dep_heading.style.font.name = 'Times New Roman'
    
    dependencies = [
        'TOTAL PARTICIPATION: Every instance MUST participate in the relationship (denoted with double line)',
        'PARTIAL PARTICIPATION: An entity MAY participate, but it is not mandatory (denoted with single line)'
    ]
    
    for item in dependencies:
        p = doc.add_paragraph(item, style='List Bullet')
        p.style.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Detailed Explanations
    rel_heading = doc.add_heading('4. DETAILED RELATIONSHIP EXPLANATIONS', level=1)
    rel_heading.style.font.name = 'Times New Roman'
    
    # CGDC_ADMIN
    cgdc_h = doc.add_heading('4.1 CGDC_ADMIN Entity', level=2)
    cgdc_h.style.font.name = 'Times New Roman'
    
    cgdc_p = doc.add_paragraph(
        'CGDC_ADMIN is the top-level administrative entity. Every CGDC admin must supervise at least one coordinator (TOTAL participation). '
        'Each coordinator must be supervised by exactly one CGDC admin (TOTAL participation from coordinator side). This creates a 1:N relationship with TOTAL participation on both sides.'
    )
    cgdc_p.style.font.name = 'Times New Roman'
    
    t1 = doc.add_table(rows=2, cols=6)
    t1.style = 'Light Grid Accent 1'
    hdr = t1.rows[0].cells
    headers = ['Relationship', 'Type', 'Cardinality', 'CGDC_ADMIN\nDependency', 'Coordinator\nDependency', 'Foreign Key']
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = 'Times New Roman'
    
    row = t1.rows[1].cells
    row_data = ['CGDC_ADMIN → PLACEMENT_COORDINATOR\n(supervises)', 'Hierarchical', '1:N', 'TOTAL', 'TOTAL', 'PLACEMENT_COORDINATOR.cgdc_id']
    for i, d in enumerate(row_data):
        row[i].text = d
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # PLACEMENT_COORDINATOR
    coord_h = doc.add_heading('4.2 PLACEMENT_COORDINATOR Entity', level=2)
    coord_h.style.font.name = 'Times New Roman'
    
    coord_p = doc.add_paragraph(
        'PLACEMENT_COORDINATOR is the middle-tier entity managing students and applications. '
        'It maintains three outgoing relationships: supervised by CGDC_ADMIN (N:1, TOTAL), coordinates STUDENT (1:N, PARTIAL), and manages APPLICATION (1:N, PARTIAL).'
    )
    coord_p.style.font.name = 'Times New Roman'
    
    t2 = doc.add_table(rows=4, cols=6)
    t2.style = 'Light Grid Accent 1'
    hdr = t2.rows[0].cells
    headers = ['Relationship', 'Type', 'Cardinality', 'Coordinator\nDependency', 'Partner\nDependency', 'Foreign Key']
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = 'Times New Roman'
    
    coord_data = [
        ['← CGDC_ADMIN (is_supervised_by)', 'Hierarchical', 'N:1', 'TOTAL', 'TOTAL', 'PLACEMENT_COORDINATOR.cgdc_id'],
        ['→ STUDENT (coordinates)', 'Supervisory', '1:N', 'PARTIAL', 'TOTAL', 'STUDENT.coord_id'],
        ['→ APPLICATION (manages)', 'Administrative', '1:N', 'PARTIAL', 'PARTIAL', 'APPLICATION.assigned_coord_id']
    ]
    
    for idx, data in enumerate(coord_data, 1):
        row = t2.rows[idx].cells
        for i, d in enumerate(data):
            row[i].text = d
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # STUDENT
    student_h = doc.add_heading('4.3 STUDENT Entity', level=2)
    student_h.style.font.name = 'Times New Roman'
    
    student_p = doc.add_paragraph(
        'STUDENT is the central entity representing students in the placement system. '
        'Each student has TOTAL participation with PLACEMENT_COORDINATOR (every student must be assigned to a coordinator). '
        'Students participate in the complete placement pipeline through APPLICATION, INTERVIEW, OFFER, and PLACEMENT_RECORD (all PARTIAL).'
    )
    student_p.style.font.name = 'Times New Roman'
    
    t3 = doc.add_table(rows=9, cols=5)
    t3.style = 'Light Grid Accent 1'
    hdr = t3.rows[0].cells
    headers = ['Relationship', 'Cardinality', 'Student\nDependency', 'Partner\nDependency', 'Foreign Key']
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = 'Times New Roman'
    
    student_data = [
        ['← PLACEMENT_COORDINATOR (coordinated_by)', 'N:1', 'TOTAL', 'PARTIAL', 'STUDENT.coord_id'],
        ['→ USER_ROLE (has_account)', '1:1', 'TOTAL', 'PARTIAL', 'USER_ROLE.entity_id'],
        ['→ RESUME (uploads)', '1:N', 'PARTIAL', 'TOTAL', 'RESUME.s_id'],
        ['→ APPLICATION (applies_to)', '1:N', 'PARTIAL', 'TOTAL', 'APPLICATION.s_id'],
        ['→ INTERVIEW (attends)', '1:N', 'PARTIAL', 'TOTAL', 'INTERVIEW.s_id'],
        ['→ OFFER (receives)', '1:N', 'PARTIAL', 'TOTAL', 'OFFER.s_id'],
        ['→ PLACEMENT_RECORD (secures)', '1:N', 'PARTIAL', 'TOTAL', 'PLACEMENT_RECORD.s_id'],
        ['→ STUDENT_SKILL (possesses)', '1:N', 'PARTIAL', 'TOTAL', 'STUDENT_SKILL.s_id']
    ]
    
    for idx, data in enumerate(student_data, 1):
        row = t3.rows[idx].cells
        for i, d in enumerate(data):
            row[i].text = d
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # COMPANY
    company_h = doc.add_heading('4.4 COMPANY Entity', level=2)
    company_h.style.font.name = 'Times New Roman'
    
    company_p = doc.add_paragraph(
        'COMPANY represents recruiting organizations. Each company posts multiple job profiles, hires multiple students through PLACEMENT_RECORD, and records campus visits through COMPANY_VISIT_HISTORY. All relationships have PARTIAL participation from the company side.'
    )
    company_p.style.font.name = 'Times New Roman'
    
    t4 = doc.add_table(rows=4, cols=5)
    t4.style = 'Light Grid Accent 1'
    hdr = t4.rows[0].cells
    headers = ['Relationship', 'Cardinality', 'Company\nDependency', 'Partner\nDependency', 'Foreign Key']
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = 'Times New Roman'
    
    company_data = [
        ['→ JOB_PROFILE (posts)', '1:N', 'PARTIAL', 'TOTAL', 'JOB_PROFILE.comp_id'],
        ['→ PLACEMENT_RECORD (hires)', '1:N', 'PARTIAL', 'TOTAL', 'PLACEMENT_RECORD.comp_id'],
        ['→ COMPANY_VISIT_HISTORY (visits)', '1:N', 'PARTIAL', 'TOTAL', 'COMPANY_VISIT_HISTORY.comp_id']
    ]
    
    for idx, data in enumerate(company_data, 1):
        row = t4.rows[idx].cells
        for i, d in enumerate(data):
            row[i].text = d
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Placement Pipeline
    pipeline_h = doc.add_heading('4.5 Placement Pipeline Relationships', level=2)
    pipeline_h.style.font.name = 'Times New Roman'
    
    pipeline_p = doc.add_paragraph(
        'The core placement process involves four entities forming the transactional pipeline: APPLICATION → INTERVIEW → OFFER → PLACEMENT_RECORD. '
        'This pipeline tracks each student\'s journey through the hiring process with PARTIAL participation (not all students apply, not all are selected).'
    )
    pipeline_p.style.font.name = 'Times New Roman'
    
    t5 = doc.add_table(rows=5, cols=4)
    t5.style = 'Light Grid Accent 1'
    hdr = t5.rows[0].cells
    headers = ['Pipeline Entity', 'Purpose', 'Key Columns', 'Cardinality']
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = 'Times New Roman'
    
    pipeline_data = [
        ['APPLICATION (Bridge)', 'Links STUDENT to JOB_PROFILE', 's_id, job_id (M:N)', 'M:N'],
        ['INTERVIEW', 'Tracks interview scheduling', 's_id, job_id', '1:N'],
        ['OFFER', 'Records job offers', 's_id, job_id', '1:N'],
        ['PLACEMENT_RECORD', 'Confirms final placement', 's_id, comp_id, job_id', '1:N']
    ]
    
    for idx, data in enumerate(pipeline_data, 1):
        row = t5.rows[idx].cells
        for i, d in enumerate(data):
            row[i].text = d
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Bridge Tables
    bridge_h = doc.add_heading('4.6 Normalized Bridge Tables (1NF Compliance)', level=2)
    bridge_h.style.font.name = 'Times New Roman'
    
    bridge_p = doc.add_paragraph(
        'To achieve strict 1NF compliance, the schema includes bridge tables that store atomic values for attributes that could contain multiple values. '
        'These normalized mappings prevent data redundancy and enable efficient querying.'
    )
    bridge_p.style.font.name = 'Times New Roman'
    
    t6 = doc.add_table(rows=5, cols=4)
    t6.style = 'Light Grid Accent 1'
    hdr = t6.rows[0].cells
    headers = ['Bridge Table', 'Purpose', 'Composite Key', 'Cardinality']
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = 'Times New Roman'
    
    bridge_data = [
        ['JOB_REQUIRED_SKILL', 'Normalizes multi-valued required skills', '(job_id, skill_name)', 'M:N'],
        ['JOB_ELIGIBILITY_BRANCH', 'Normalizes eligible departments for jobs', '(job_id, branch_name)', 'M:N'],
        ['VISIT_COVERED_STREAM', 'Maps departments covered in company visits', '(visit_id, stream_name)', 'M:N'],
        ['RESUME_PARSED_KEYWORD', 'Normalizes extracted keywords from resumes', '(resume_id, keyword)', 'M:N']
    ]
    
    for idx, data in enumerate(bridge_data, 1):
        row = t6.rows[idx].cells
        for i, d in enumerate(data):
            row[i].text = d
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Polymorphic Design
    poly_h = doc.add_heading('4.7 Polymorphic Authentication (USER_ROLE)', level=2)
    poly_h.style.font.name = 'Times New Roman'
    
    poly_p = doc.add_paragraph(
        'The USER_ROLE entity implements polymorphic authentication to serve three different user types with a single login table. '
        'The entity_id column polymorphically references s_id, coord_id, or cgdc_id depending on the role field. '
        'Each user type has a 1:1 relationship with USER_ROLE (TOTAL participation from user side, PARTIAL from USER_ROLE side).'
    )
    poly_p.style.font.name = 'Times New Roman'
    
    t7 = doc.add_table(rows=4, cols=3)
    t7.style = 'Light Grid Accent 1'
    hdr = t7.rows[0].cells
    headers = ['User Role', 'entity_id Reference', 'Cardinality']
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = 'Times New Roman'
    
    poly_data = [
        ['student', 'STUDENT.s_id', '1:1'],
        ['coordinator', 'PLACEMENT_COORDINATOR.coord_id', '1:1'],
        ['cgdc_admin', 'CGDC_ADMIN.cgdc_id', '1:1']
    ]
    
    for idx, data in enumerate(poly_data, 1):
        row = t7.rows[idx].cells
        for i, d in enumerate(data):
            row[i].text = d
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Key Design Principles
    principles_h = doc.add_heading('5. KEY DESIGN PRINCIPLES', level=1)
    principles_h.style.font.name = 'Times New Roman'
    
    principles = [
        'Hierarchical Authority: Three-tier structure (Admin → Coordinator → Student) ensures proper oversight',
        'Referential Integrity: Foreign keys with CASCADE delete maintain data consistency and prevent orphaned records',
        'Normalization: Bridge tables and atomic values achieve full 1NF/2NF/3NF compliance',
        'Polymorphism: Single USER_ROLE table serves three user types, reducing schema complexity',
        'Audit Trail: STATUS_AUDIT_LOG is auto-populated by database triggers for compliance tracking',
        'M:N Resolution: APPLICATION entity elegantly resolves many-to-many relationship between STUDENT and JOB_PROFILE',
        'Transactional Integrity: Complete placement pipeline supports ACID transactions for data reliability'
    ]
    
    for principle in principles:
        p = doc.add_paragraph(principle, style='List Bullet')
        p.style.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Conclusion
    conclusion_h = doc.add_heading('6. CONCLUSION', level=1)
    conclusion_h.style.font.name = 'Times New Roman'
    
    conclusion_p = doc.add_paragraph(
        'The ER diagram represents a meticulously designed, normalized database schema that comprehensively captures the student placement ecosystem. '
        'With 22 entities, 40+ relationships, and precise cardinality/dependency definitions, the system ensures data consistency, referential integrity, and supports ACID transactions. '
        'The hierarchical structure provides clear lines of authority, while normalized bridge tables eliminate redundancy and achieve 3NF compliance. '
        'The polymorphic authentication design and comprehensive audit trail demonstrate professional-grade database engineering suitable for production deployment.'
    )
    conclusion_p.style.font.name = 'Times New Roman'
    
    # Save
    doc.save('ER_DIAGRAM_EXPLANATION.docx')
    print("✓ Document created successfully: ER_DIAGRAM_EXPLANATION.docx")

if __name__ == '__main__':
    create_er_explanation()
